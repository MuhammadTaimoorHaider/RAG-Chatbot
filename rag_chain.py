"""
Core RAG (Retrieval Augmented Generation) logic.
Handles document processing, embeddings, vector storage, and conversational retrieval.
"""

import os
import json
import time
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

# Document Loading
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Groq & Embeddings
from langchain_groq import ChatGroq
from langchain_community.embeddings import HuggingFaceEmbeddings

# Pinecone
from pinecone import Pinecone, ServerlessSpec
from langchain_community.vectorstores import Pinecone as LangchainPinecone

# LangChain
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.schema import BaseMessage, HumanMessage, AIMessage
from langchain.schema.chat_history import BaseChatMessageHistory

# Redis for conversation memory
import redis

# Document processing
import docx
import markdown

# Logging
from loguru import logger

# Configuration
from config import settings


# Custom Exceptions
class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


class VectorStoreError(Exception):
    """Raised when vector store operations fail."""
    pass


class LLMError(Exception):
    """Raised when LLM generation fails."""
    pass


class DocumentProcessor:
    """Handles document loading, parsing, and chunking."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        logger.info(f"DocumentProcessor initialized with chunk_size={settings.chunk_size}, overlap={settings.chunk_overlap}")

    def load_pdf(self, file_path: str) -> List[Document]:
        """Extract text from PDF file."""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            logger.info(f"Loaded PDF: {file_path} ({len(documents)} pages)")
            return documents
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {str(e)}")
            raise DocumentProcessingError(f"PDF loading failed: {str(e)}")

    def load_docx(self, file_path: str) -> List[Document]:
        """Extract text from Word document."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

            document = Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "type": "docx"
                }
            )
            logger.info(f"Loaded DOCX: {file_path} ({len(text)} characters)")
            return [document]
        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {str(e)}")
            raise DocumentProcessingError(f"DOCX loading failed: {str(e)}")

    def load_text(self, file_path: str) -> List[Document]:
        """Load plain text or markdown file."""
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()

            # Add metadata
            for doc in documents:
                doc.metadata.update({
                    "filename": os.path.basename(file_path),
                    "type": "txt" if file_path.endswith('.txt') else "md"
                })

            logger.info(f"Loaded text file: {file_path}")
            return documents
        except Exception as e:
            logger.error(f"Failed to load text file {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Text loading failed: {str(e)}")

    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks with metadata."""
        try:
            chunks = self.text_splitter.split_documents(documents)

            # Add chunk metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "chunk_id": i,
                    "chunk_size": len(chunk.page_content),
                    "timestamp": datetime.utcnow().isoformat()
                })

            logger.info(f"Created {len(chunks)} chunks from {len(documents)} documents")
            return chunks
        except Exception as e:
            logger.error(f"Failed to chunk documents: {str(e)}")
            raise DocumentProcessingError(f"Chunking failed: {str(e)}")

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file."""
        file_path = Path(file_path)
        return {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size if file_path.exists() else 0,
            "file_type": file_path.suffix,
            "upload_time": datetime.utcnow().isoformat()
        }


class EmbeddingManager:
    """Manages HuggingFace embeddings generation (free, local)."""

    def __init__(self, model: Optional[str] = None):
        self.model = model or settings.embedding_model

        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.model,
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info(f"EmbeddingManager initialized with model: {self.model}")
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {str(e)}")
            raise LLMError(f"Embedding initialization failed: {str(e)}")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple documents."""
        try:
            embeddings = self.embeddings.embed_documents(texts)
            logger.debug(f"Generated embeddings for {len(texts)} documents")
            return embeddings
        except Exception as e:
            logger.error(f"Failed to embed documents: {str(e)}")
            raise LLMError(f"Document embedding failed: {str(e)}")

    def embed_query(self, text: str) -> List[float]:
        """Generate embedding for a single query."""
        try:
            embedding = self.embeddings.embed_query(text)
            logger.debug(f"Generated query embedding (dimension: {len(embedding)})")
            return embedding
        except Exception as e:
            logger.error(f"Failed to embed query: {str(e)}")
            raise LLMError(f"Query embedding failed: {str(e)}")


class VectorStoreManager:
    """Manages Pinecone vector database operations."""

    def __init__(
        self,
        api_key: Optional[str] = None,
        environment: Optional[str] = None,
        index_name: Optional[str] = None,
        embedding_dimension: int = 1536
    ):
        self.api_key = api_key or settings.pinecone_api_key
        self.environment = environment or settings.pinecone_environment
        self.index_name = index_name or settings.pinecone_index_name
        self.embedding_dimension = embedding_dimension

        try:
            self.pc = Pinecone(api_key=self.api_key)
            logger.info(f"Pinecone client initialized (environment: {self.environment})")
        except Exception as e:
            logger.error(f"Failed to initialize Pinecone: {str(e)}")
            raise VectorStoreError(f"Pinecone initialization failed: {str(e)}")

    def create_index(self, dimension: Optional[int] = None, metric: str = "cosine") -> bool:
        """Create Pinecone index if it doesn't exist."""
        try:
            target_dimension = dimension or self.embedding_dimension
            existing_indexes = [index.name for index in self.pc.list_indexes()]

            if self.index_name not in existing_indexes:
                self.pc.create_index(
                    name=self.index_name,
                    dimension=target_dimension,
                    metric=metric,
                    spec=ServerlessSpec(
                        cloud='aws',
                        region=self.environment
                    )
                )
                logger.info(f"Created Pinecone index: {self.index_name}")
                return True
            else:
                # If the configured index exists with an incompatible dimension,
                # fall back to a dimension-specific index to keep runtime working.
                index_info = self.pc.describe_index(self.index_name)
                existing_dimension = getattr(index_info, "dimension", None)

                if existing_dimension is not None and existing_dimension != target_dimension:
                    fallback_index_name = f"{self.index_name}-{target_dimension}d"
                    logger.warning(
                        f"Index '{self.index_name}' has dimension {existing_dimension}, "
                        f"but embeddings are {target_dimension}. Using '{fallback_index_name}' instead."
                    )
                    self.index_name = fallback_index_name

                    existing_indexes = [index.name for index in self.pc.list_indexes()]
                    if self.index_name not in existing_indexes:
                        self.pc.create_index(
                            name=self.index_name,
                            dimension=target_dimension,
                            metric=metric,
                            spec=ServerlessSpec(
                                cloud='aws',
                                region=self.environment
                            )
                        )
                        logger.info(f"Created Pinecone fallback index: {self.index_name}")
                        return True

                logger.info(f"Pinecone index already exists: {self.index_name}")
                return False
        except Exception as e:
            logger.error(f"Failed to create index: {str(e)}")
            raise VectorStoreError(f"Index creation failed: {str(e)}")

    def get_vectorstore(self, embeddings: EmbeddingManager, namespace: str = "default") -> LangchainPinecone:
        """Get LangChain Pinecone vectorstore instance."""
        try:
            index = self.pc.Index(self.index_name)
            vectorstore = LangchainPinecone(
                index=index,
                embedding=embeddings.embeddings,
                text_key="text",
                namespace=namespace
            )
            logger.info(f"Retrieved vectorstore for index: {self.index_name}, namespace: {namespace}")
            return vectorstore
        except Exception as e:
            logger.error(f"Failed to get vectorstore: {str(e)}")
            raise VectorStoreError(f"Vectorstore retrieval failed: {str(e)}")

    def upsert_documents(self, documents: List[Document], embeddings: EmbeddingManager, namespace: str = "default") -> Dict[str, Any]:
        """Store document embeddings in Pinecone."""
        try:
            start_time = time.time()

            vectorstore = self.get_vectorstore(embeddings, namespace)
            vectorstore.add_documents(documents)

            elapsed_time = time.time() - start_time

            logger.info(f"Upserted {len(documents)} documents to namespace '{namespace}' in {elapsed_time:.2f}s")
            return {
                "status": "success",
                "documents_processed": len(documents),
                "namespace": namespace,
                "elapsed_time": elapsed_time
            }
        except Exception as e:
            logger.error(f"Failed to upsert documents: {str(e)}")
            raise VectorStoreError(f"Document upsert failed: {str(e)}")

    def similarity_search(self, query: str, embeddings: EmbeddingManager, top_k: int = 4, namespace: str = "default") -> List[Document]:
        """Perform similarity search in Pinecone."""
        try:
            vectorstore = self.get_vectorstore(embeddings, namespace)
            results = vectorstore.similarity_search(query, k=top_k)

            logger.debug(f"Similarity search returned {len(results)} results")
            return results
        except Exception as e:
            logger.error(f"Failed to perform similarity search: {str(e)}")
            raise VectorStoreError(f"Similarity search failed: {str(e)}")

    def delete_namespace(self, namespace: str):
        """Delete all vectors in a namespace."""
        try:
            index = self.pc.Index(self.index_name)
            index.delete(delete_all=True, namespace=namespace)
            logger.info(f"Deleted namespace: {namespace}")
        except Exception as e:
            logger.error(f"Failed to delete namespace: {str(e)}")
            raise VectorStoreError(f"Namespace deletion failed: {str(e)}")


class ConversationManager:
    """Manages conversation history with hybrid memory (Redis + in-memory)."""

    def __init__(self, max_history: int = None):
        self.max_history = max_history or settings.max_conversation_history
        self.redis_available = self._check_redis()

        if self.redis_available:
            try:
                self.redis_client = redis.Redis(
                    host=settings.redis_host,
                    port=settings.redis_port,
                    db=settings.redis_db,
                    password=settings.redis_password,
                    decode_responses=True
                )
                logger.info("ConversationManager using Redis for memory")
            except Exception as e:
                logger.warning(f"Redis connection failed, falling back to in-memory: {str(e)}")
                self.redis_available = False
                self.memory = {}
        else:
            self.memory = {}
            logger.info("ConversationManager using in-memory storage")

    def _check_redis(self) -> bool:
        """Check if Redis is available."""
        try:
            r = redis.Redis(
                host=settings.redis_host,
                port=settings.redis_port,
                db=settings.redis_db,
                password=settings.redis_password,
                socket_connect_timeout=2
            )
            r.ping()
            return True
        except:
            return False

    def add_message(self, session_id: str, role: str, content: str):
        """Add a message to conversation history."""
        message = {
            "role": role,
            "content": content,
            "timestamp": datetime.utcnow().isoformat()
        }

        if self.redis_available:
            try:
                key = f"conversation:{session_id}"
                self.redis_client.lpush(key, json.dumps(message))
                self.redis_client.ltrim(key, 0, 49)  # Keep last 50 messages
                self.redis_client.expire(key, 3600)  # 1 hour TTL
            except Exception as e:
                logger.error(f"Redis error, falling back to in-memory: {str(e)}")
                self._add_to_memory(session_id, message)
        else:
            self._add_to_memory(session_id, message)

    def _add_to_memory(self, session_id: str, message: Dict):
        """Add message to in-memory storage."""
        if session_id not in self.memory:
            self.memory[session_id] = []
        self.memory[session_id].append(message)
        # Keep last N messages
        self.memory[session_id] = self.memory[session_id][-self.max_history * 2:]

    def get_history(self, session_id: str, limit: int = 5) -> List[Dict]:
        """Retrieve conversation history."""
        if self.redis_available:
            try:
                key = f"conversation:{session_id}"
                messages = self.redis_client.lrange(key, 0, limit - 1)
                return [json.loads(msg) for msg in reversed(messages)]
            except Exception as e:
                logger.error(f"Redis error, falling back to in-memory: {str(e)}")
                return self.memory.get(session_id, [])[-limit:]
        else:
            return self.memory.get(session_id, [])[-limit:]

    def clear_history(self, session_id: str):
        """Clear conversation history for a session."""
        if self.redis_available:
            try:
                key = f"conversation:{session_id}"
                self.redis_client.delete(key)
                logger.info(f"Cleared Redis history for session: {session_id}")
            except Exception as e:
                logger.error(f"Redis error: {str(e)}")

        if session_id in self.memory:
            del self.memory[session_id]
            logger.info(f"Cleared in-memory history for session: {session_id}")

    def format_history_for_prompt(self, session_id: str, limit: int = 5) -> str:
        """Format conversation history for LLM prompt."""
        history = self.get_history(session_id, limit)
        if not history:
            return ""

        formatted = []
        for msg in history:
            role = "Human" if msg["role"] == "user" else "Assistant"
            formatted.append(f"{role}: {msg['content']}")

        return "\n".join(formatted)


class CustomChatHistory(BaseChatMessageHistory):
    """Custom chat history for LangChain backed by ConversationManager."""

    def __init__(self, session_id: str, conversation_manager: ConversationManager):
        self.session_id = session_id
        self.conversation_manager = conversation_manager

    def add_message(self, message: BaseMessage):
        """Add a message to the history."""
        role = "user" if isinstance(message, HumanMessage) else "assistant"
        self.conversation_manager.add_message(
            self.session_id, role, message.content
        )

    def clear(self):
        """Clear the conversation history."""
        self.conversation_manager.clear_history(self.session_id)

    @property
    def messages(self) -> List[BaseMessage]:
        """Retrieve messages as LangChain BaseMessage objects."""
        history = self.conversation_manager.get_history(self.session_id)
        return [
            HumanMessage(content=msg["content"])
            if msg["role"] == "user"
            else AIMessage(content=msg["content"])
            for msg in history
        ]


class RAGChain:
    """Main RAG chain orchestrator."""

    def __init__(
        self,
        vector_store_manager: VectorStoreManager,
        embedding_manager: EmbeddingManager,
        conversation_manager: ConversationManager,
        llm_api_key: Optional[str] = None,
        llm_model: Optional[str] = None
    ):
        self.vector_store_manager = vector_store_manager
        self.embedding_manager = embedding_manager
        self.conversation_manager = conversation_manager

        self.llm_api_key = llm_api_key or settings.groq_api_key
        self.llm_model = llm_model or settings.groq_model
        self._fallback_models = [
            "llama-3.3-70b-versatile",
            "llama-3.1-8b-instant"
        ]

        # Initialize Groq LLM
        self.llm = ChatGroq(
            model=self.llm_model,
            groq_api_key=self.llm_api_key,
            temperature=0
        )

        logger.info(f"RAGChain initialized with model: {self.llm_model}")

    def _activate_fallback_model(self) -> bool:
        """Switch to the next available Groq model if current model is unavailable."""
        for model_name in self._fallback_models:
            if model_name == self.llm_model:
                continue

            try:
                self.llm = ChatGroq(
                    model=model_name,
                    groq_api_key=self.llm_api_key,
                    temperature=0
                )
                self.llm_model = model_name
                logger.warning(f"Switched to fallback Groq model: {self.llm_model}")
                return True
            except Exception as e:
                logger.warning(f"Failed to initialize fallback model {model_name}: {str(e)}")

        return False

    def get_prompt_template(self) -> PromptTemplate:
        """Get the prompt template for RAG."""
        template = """You are a helpful AI assistant that answers questions based on the provided context.

Context from documents:
{context}

Conversation history:
{chat_history}

Current question: {question}

Instructions:
- Answer based on the context provided
- If the answer is not in the context, say "I don't have enough information to answer that"
- Cite sources when possible (mention document names or page numbers)
- Be concise but thorough

Answer:"""

        return PromptTemplate(
            template=template,
            input_variables=["context", "chat_history", "question"]
        )

    def create_chain(self, namespace: str = "default") -> ConversationalRetrievalChain:
        """Create a conversational retrieval chain."""
        try:
            # Get vectorstore
            vectorstore = self.vector_store_manager.get_vectorstore(
                self.embedding_manager,
                namespace
            )

            # Create retriever
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={"k": settings.top_k_results}
            )

            # Create chain
            chain = ConversationalRetrievalChain.from_llm(
                llm=self.llm,
                retriever=retriever,
                return_source_documents=True,
                verbose=False
            )

            return chain
        except Exception as e:
            logger.error(f"Failed to create chain: {str(e)}")
            raise LLMError(f"Chain creation failed: {str(e)}")

    def query(
        self,
        question: str,
        session_id: str,
        namespace: str = "default"
    ) -> Dict[str, Any]:
        """Query the RAG system with conversation memory."""
        try:
            start_time = time.time()

            # Get conversation history
            chat_history = self.conversation_manager.get_history(session_id, limit=5)
            formatted_history = [(msg["content"] if msg["role"] == "user" else "",
                                 msg["content"] if msg["role"] == "assistant" else "")
                                for msg in chat_history]

            # Create chain
            chain = self.create_chain(namespace)

            # Query
            retrieval_start = time.time()
            try:
                response = chain.invoke({
                    "question": question,
                    "chat_history": formatted_history
                })
            except Exception as model_error:
                error_text = str(model_error).lower()
                if "model_decommissioned" in error_text or "decommissioned" in error_text:
                    logger.warning("Configured Groq model is decommissioned. Attempting fallback model.")
                    if self._activate_fallback_model():
                        chain = self.create_chain(namespace)
                        response = chain.invoke({
                            "question": question,
                            "chat_history": formatted_history
                        })
                    else:
                        raise
                else:
                    raise
            retrieval_time = time.time() - retrieval_start

            # Extract answer and sources
            answer = response.get("answer", "")
            source_documents = response.get("source_documents", [])

            # Format sources
            sources = []
            for i, doc in enumerate(source_documents):
                sources.append({
                    "id": i,
                    "content": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                    "metadata": doc.metadata,
                    "filename": doc.metadata.get("filename", "unknown"),
                    "page": doc.metadata.get("page", 0),
                    "chunk_id": doc.metadata.get("chunk_id", 0)
                })

            # Save to conversation history
            self.conversation_manager.add_message(session_id, "user", question)
            self.conversation_manager.add_message(session_id, "assistant", answer)

            elapsed_time = time.time() - start_time

            result = {
                "answer": answer,
                "sources": sources,
                "session_id": session_id,
                "metadata": {
                    "model": self.llm_model,
                    "retrieval_time_ms": int(retrieval_time * 1000),
                    "total_time_ms": int(elapsed_time * 1000),
                    "num_sources": len(sources),
                    "namespace": namespace
                }
            }

            logger.info(f"Query processed in {elapsed_time:.2f}s for session {session_id}")
            return result

        except Exception as e:
            logger.error(f"Query failed: {str(e)}")
            raise LLMError(f"Query processing failed: {str(e)}")


# Initialize global instances (for use in FastAPI)
def initialize_rag_system():
    """Initialize core RAG components."""
    try:
        embedding_manager = EmbeddingManager()

        # Keep Pinecone index dimension aligned with the configured embedding model.
        embedding_dimension = len(embedding_manager.embed_query("dimension probe"))
        vector_store_manager = VectorStoreManager(embedding_dimension=embedding_dimension)
        conversation_manager = ConversationManager()

        # Ensure index exists
        vector_store_manager.create_index()

        rag_chain = RAGChain(
            vector_store_manager=vector_store_manager,
            embedding_manager=embedding_manager,
            conversation_manager=conversation_manager
        )

        logger.info("RAG system initialized successfully")
        return {
            "embedding_manager": embedding_manager,
            "vector_store_manager": vector_store_manager,
            "conversation_manager": conversation_manager,
            "rag_chain": rag_chain,
            "document_processor": DocumentProcessor()
        }
    except Exception as e:
        logger.error(f"Failed to initialize RAG system: {str(e)}")
        raise RuntimeError(f"RAG system initialization failed: {str(e)}")
